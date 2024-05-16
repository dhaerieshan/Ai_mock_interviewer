import openai
from docx import Document

class EvaluationService:
    def __init__(self, api_key, api_version, azure_endpoint):
        self.client = openai.AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=api_key,
            api_version=api_version
        )

    def evaluate_interview(self, qa_pairs):
        interview_content = ""
        for question, answer in qa_pairs:
            interview_content += f"Question: {question}\nAnswer: {answer}\n\n"

        # First, we get general feedback and scores for the interview.
        general_feedback = self.generate_general_feedback(interview_content)
        
        # Then, evaluate speaking skills and tone.
        speaking_skills_feedback = self.evaluate_speaking_skills(interview_content)
        tone_feedback = self.evaluate_tone(interview_content)

        # Combine all parts of the evaluation into one document.
        full_evaluation = f"{general_feedback}\n\nSpeaking Skills Feedback:\n{speaking_skills_feedback}\n\nTone Feedback:\n{tone_feedback}"
        
        return full_evaluation

    def generate_general_feedback(self, interview_content):
        system_message = """You are a training tool simulating a strict interview environment. Evaluate the entire interview content by considering its relevance, conciseness, and how well the candidate has communicated their qualifications and suitability for the job.
        - Summarization: Provide an overall summary of the candidate's performance.
        - Strengths: Highlight the positive aspects throughout the interview.
        - Areas for Improvement: Suggest areas where the candidate could improve their responses.
        - Score out of 100: Give a numerical score based on the overall quality and relevance of the interview.
        - Actionable Tips: Provide actionable tips to help the candidate improve their interview skills.
        - Sample Answers: Offer examples of ideal responses for future reference."""

        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "system", "content": system_message}, {"role": "user", "content": interview_content}],
            temperature=0.7,
            max_tokens=10000,
            top_p=1.0,
            frequency_penalty=0.0,
            presence_penalty=0.0
        )
        return response.choices[0].message.content

    def evaluate_speaking_skills(self, interview_content):
        prompt = f"""Based on the transcribed text of the interview responses, evaluate the candidate's English speaking skills focusing on grammar, vocabulary, and fluency. Note: Pronunciation and other audio-specific characteristics cannot be assessed from the text and require direct audio analysis:{interview_content}"""
        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "system", "content": prompt}],
            temperature=0.5,
            max_tokens=10000,
            top_p=1.0
        )
        return response.choices[0].message.content

    def evaluate_tone(self, interview_content):
        prompt = f"Analyze the emotional tone of the following interview responses. Provide insights on the candidate's confidence, enthusiasm, and overall communication effectiveness: {interview_content}"
        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "system", "content": prompt}],
            temperature=0.5,
            max_tokens=10000,
            top_p=1.0
        )
        return response.choices[0].message.content

    def save_evaluation_to_doc(self, evaluation, filename="evaluation.docx"):
        # Create a new Document
        doc = Document()
        
        # Add a heading
        doc.add_heading('Interview Evaluation', level=1)

        # Add the evaluation content to the document
        doc.add_paragraph(evaluation)

        # Save the document
        doc.save(filename)
        print(f"Evaluation saved to {filename}")
