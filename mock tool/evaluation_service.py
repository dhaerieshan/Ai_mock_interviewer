import openai
from docx import Document


class EvaluationService:
    def __init__(self, api_key, api_version, azure_endpoint):
        self.client = openai.AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=api_key,
            api_version=api_version
        )

    def evaluate_interview(self, qa_pairs):
        interview_content = ""
        for question, answer in qa_pairs:
            interview_content += f"Question: {question}\nAnswer: {answer}\n\n"

        system_message = """You are a training tool simulating a strict interview environment. Evaluate the entire interview content by considering its relevance, conciseness, and how well the candidate has communicated their qualifications and suitability for the job.
        - Summarization: Provide an overall summary of the candidate's performance.
        - Strengths: Highlight the positive aspects throughout the interview.
        - Areas for Improvement: Suggest areas where the candidate could improve their responses.
        - Score out of 100: Give a numerical score based on the overall quality and relevance of the interview.
        - Actionable Tips: Provide actionable tips to help the candidate improve their interview skills.
        - Sample Answers: Offer examples of ideal responses for future reference."""

        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": interview_content}
        ]

        response = self.client.chat_completions.create(
            model="gpt-4",  # Make sure to use your actual model identifier
            messages=messages,
            temperature=0.7,
            max_tokens=1500,
            top_p=1.0,
            frequency_penalty=0.0,
            presence_penalty=0.0
        )

        evaluation = response.choices[0].message.content.strip()
        return evaluation

    def save_evaluation_to_doc(self, evaluation, filename="evaluation.docx"):
        # Create a new Document
        doc = Document()
        
        # Add a heading
        doc.add_heading('Interview Evaluation', level=1)

        # Add the evaluation content to the document
        doc.add_paragraph(evaluation)

        # Save the document
        doc.save(filename)
        print(f"Evaluation saved to {filename}")